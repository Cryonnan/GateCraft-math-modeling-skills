# -*- coding: utf-8 -*-
"""
图片/纯图像 PDF -> Word（OCR 重建流水线，数学公式多、图像多的场景）

适用场景:
    - PDF 无文本层（扫描件/课件截图），或图片型 PDF
    - 数学公式多、需要 Word 原生可编辑公式（OMML）
    - 图片需要保留原图嵌入

流程:
    1. 渲染 PDF 每页为 PNG（dpi=150）
    2. Qwen3-VL 逐页 OCR -> Markdown（公式 LaTeX，图片占位 [图N]）
    3. 提取 PDF 内嵌图片（命名 p{i}_img{j}.png）
    4. 组装 combined.md：把 [图N] 占位符替换为 ![](imgs/...){width=...}
       （IMG_MAP 配置 页号->图号->文件名+宽度；未配置的占位符保留原文）
    5. 清洗 markdown：修复 OCR 常见问题（$ 前空格、全角符号等）
    6. pandoc --reference-doc -> docx（公式转 OMML 可编辑）
    7. 验证输出 docx 结构统计

用法:
    python pdf_image_to_docx.py input.pdf [--ref ref.docx|ref.dotx] [--out out.docx]

依赖: pandoc, python-docx, PyMuPDF(fitz), openai, 网络（SiliconFlow API）
"""

import argparse
import glob
import os
import re
import subprocess
import sys

if sys.stdout and sys.stdout.encoding and sys.stdout.encoding.lower() not in ("utf-8", "utf8"):
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

HERE = os.path.dirname(os.path.abspath(__file__))

# ============ 配置区（按需修改） ============
# (页码, 图序号) -> (文件名, 显示宽度)。示例：
#   IMG_MAP = {(3,1): ('p3_img3.png', '9cm'), (3,2): ('p3_img6.png', '12cm')}
# 自动推导：若文件在 imgs/ 下且键对应页内图序号，可自动填（见 build 中 auto_map 逻辑）
IMG_MAP = {}

DEFAULT_REF = os.path.join(HERE, "ref_math_model.docx")
# ===========================================


def render_pdf(pdf_path, out_dir, dpi=150):
    import fitz
    os.makedirs(out_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    paths = []
    for i, page in enumerate(doc):
        p = os.path.join(out_dir, f"page{i+1}.png")
        page.get_pixmap(dpi=dpi).save(p)
        paths.append(p)
    return paths


def extract_imgs(pdf_path, out_dir):
    import fitz
    os.makedirs(out_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    files = []
    for pno in range(len(doc)):
        for j, img in enumerate(doc[pno].get_images(full=True), 1):
            xref = img[0]
            pix = fitz.Pixmap(doc, xref)
            if pix.n - pix.alpha >= 4:
                pix = fitz.Pixmap(fitz.csRGB, pix)
            name = f"p{pno+1}_img{j}.png"
            path = os.path.join(out_dir, name)
            pix.save(path)
            files.append(path)
    return files


def auto_map_imgs(ocr_dir, imgs_dir):
    """启发式自动建图: 每页 [图N] 占位按顺序对应 imgs/p{page}_img{N}.png
    返回 {(page, n): (fname, width_cm)}；宽度按图片像素换算（按 96dpi 屏幕基准）。"""
    from PIL import Image
    map_ = {}
    ocr_files = sorted(glob.glob(os.path.join(ocr_dir, "page*.md")),
                       key=lambda p: int(re.search(r'page(\d+)\.md', p).group(1)))
    for md_path in ocr_files:
        pno = int(re.search(r'page(\d+)\.md', md_path).group(1))
        md = open(md_path, encoding="utf-8").read()
        n_places = len(re.findall(r'\[图(\d+)\]', md))
        for n in range(1, n_places + 1):
            fname = f"p{pno}_img{n}.png"
            fpath = os.path.join(imgs_dir, fname)
            if not os.path.exists(fpath):
                # 容忍 页内序号 与 内嵌图序号 错位：往后找未用的
                cands = sorted(glob.glob(os.path.join(imgs_dir, f"p{pno}_img*.png")))
                used = {v[0] for v in map_.values()}
                cands = [c for c in cands if os.path.basename(c) not in used]
                if not cands:
                    continue
                fpath = cands[0]
                fname = os.path.basename(fpath)
            try:
                w, h = Image.open(fpath).size
                width = f"{min(13, round(w * 2.54 / 96, 1))}cm"
            except Exception:
                width = "10cm"
            map_[(pno, n)] = (fname, width)
    return map_


def clean_md(md):
    """修复 OCR 输出中常见的 pandoc/数学 问题"""
    # 1. 美元前空格 -> 数学结束符必须紧贴内容（pandoc tex_math_dollars 规则）
    md = re.sub(r'[ \t]+\$', '$', md)
    # 2. 多余全角空格
    md = md.replace('\u3000', ' ')
    # 3. 公式行内连续空行清理（可加更多规则）
    return md


def build_combined(ocr_dir, img_map, out_md):
    parts = []
    ocr_files = sorted(glob.glob(os.path.join(ocr_dir, "page*.md")),
                       key=lambda p: int(re.search(r'page(\d+)\.md', p).group(1)))
    for md_path in ocr_files:
        pno = int(re.search(r'page(\d+)\.md', md_path).group(1))
        md = open(md_path, encoding="utf-8").read()

        def repl(m):
            n = int(m.group(1))
            key = (pno, n)
            if key in img_map:
                fname, w = img_map[key]
                return f'\n\n![](imgs/{fname}){{width={w}}}\n\n'
            return m.group(0)

        md = re.sub(r'\[图(\d+)\]（[^）]*）', repl, md)
        md = re.sub(r'\[图(\d+)\]\([^)]*\)', repl, md)
        parts.append(clean_md(md))
    combined = '\n\n---\n\n'.join(parts)
    with open(out_md, "w", encoding="utf-8") as f:
        f.write(combined)
    return combined


def verify_docx(docx_path):
    import docx as _docx
    from docx.oxml.ns import qn
    d = _docx.Document(docx_path)
    body = d.element.body
    n_img = len(body.findall(".//" + qn("w:drawing")))
    n_omml = len(body.findall(".//" + qn("m:oMath")))
    print(f"验证: 段落 {len(d.paragraphs)}, 表格 {len(d.tables)}, "
          f"图片 {n_img}, OMML公式 {n_omml}")
    return n_img, n_omml


def main():
    ap = argparse.ArgumentParser(description="图片/纯图像 PDF -> Word（OCR 重建）")
    ap.add_argument("pdf", help="输入 PDF")
    ap.add_argument("--ref", default="", help="reference docx/dotx（默认脚本目录 ref_math_model.docx）")
    ap.add_argument("--out", default="", help="输出 docx（默认 input.docx）")
    ap.add_argument("--ocr-dir", default="", help="OCR 输出目录（默认 pdf同目录/ocr）")
    ap.add_argument("--imgs-dir", default="", help="图片输出目录（默认 pdf同目录/imgs）")
    ap.add_argument("--no-ocr", action="store_true", help="跳过 OCR，只用已有 ocr/*.md")
    ap.add_argument("--dpi", type=int, default=150)
    ap.add_argument("--img-map", default="", help="手动图映射 JSON，如 '{(3,1):[\"p3_img3.png\",\"9cm\"]}'")
    args = ap.parse_args()

    pdf = os.path.abspath(args.pdf)
    base = os.path.splitext(pdf)[0]
    workdir = os.path.dirname(pdf)
    ref = os.path.abspath(args.ref) if args.ref else (
        DEFAULT_REF if os.path.exists(DEFAULT_REF) else "")
    out_docx = os.path.abspath(args.out) if args.out else base + ".docx"
    ocr_dir = os.path.abspath(args.ocr_dir) if args.ocr_dir else os.path.join(workdir, "ocr")
    imgs_dir = os.path.abspath(args.imgs_dir) if args.imgs_dir else os.path.join(workdir, "imgs")

    sys.path.insert(0, HERE)
    from vision_qwen import ocr_image, MATH_PROMPT, encode_image

    # 1. 渲染
    print("[1/6] 渲染 PDF 页面...")
    render_dir = os.path.join(ocr_dir, "_pages")
    pages = render_pdf(pdf, render_dir, dpi=args.dpi)
    print(f"      {len(pages)} 页 -> {render_dir}")

    # 2. OCR
    if not args.no_ocr:
        print("[2/6] Qwen3-VL 逐页 OCR...")
        os.makedirs(ocr_dir, exist_ok=True)
        for i, p in enumerate(pages, 1):
            md_path = os.path.join(ocr_dir, f"page{i}.md")
            txt = ocr_image(p, prompt=MATH_PROMPT, max_tokens=8000)
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(txt)
            print(f"      page{i} done, {len(txt)} chars")
    else:
        print("[2/6] 跳过 OCR（使用已有 ocr/*.md）")

    # 3. 提取内嵌图
    print("[3/6] 提取 PDF 内嵌图片...")
    extract_imgs(pdf, imgs_dir)

    # 4. 图映射
    print("[4/6] 建立 图片映射...")
    if args.img_map:
        img_map = {}
        for m in re.finditer(r'\((\d+),\s*(\d+)\)\s*:\s*\["([^"]+)",\s*"([^"]+)"\]', args.img_map):
            img_map[(int(m.group(1)), int(m.group(2)))] = (m.group(3), m.group(4))
        if not img_map:
            print("警告: --img-map 格式未匹配，示例: '(1,1):[\"p1_img1.png\",\"9cm\"]'")
    elif IMG_MAP:
        img_map = IMG_MAP
    else:
        img_map = auto_map_imgs(ocr_dir, imgs_dir)
    print(f"      映射 {len(img_map)} 张图: {img_map}")

    # 5. 组装 + pandoc
    print("[5/6] 组装 combined.md + pandoc 转换...")
    combined_md = os.path.join(workdir, "combined.md")
    build_combined(ocr_dir, img_map, combined_md)
    if not ref:
        cand = glob.glob(os.path.join(workdir, "*.dotx")) + \
               glob.glob(os.path.join(workdir, "*模板*.docx"))
        ref = cand[0] if cand else ""
    if not ref:
        print("错误: 未找到 reference 模板，请用 --ref 指定")
        sys.exit(1)
    subprocess.run(["pandoc", combined_md, "-f", "markdown", "-t", "docx",
                    "--reference-doc", ref, "-o", out_docx],
                   check=True, cwd=workdir)
    print(f"      -> {out_docx}")

    # 6. 验证
    print("[6/6] 验证...")
    verify_docx(out_docx)
    print("完成:", out_docx)


if __name__ == "__main__":
    main()