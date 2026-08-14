# -*- coding: utf-8 -*-
"""
PDF/TeX -> Word 转换器（数模论文专用，稳健流水线）

用法:
    python convert_to_docx.py input.tex [reference.docx|dotx] [output.docx]

流程:
    1. 图片修复: tex 引用的图片缺失但同目录有同名 PDF 时, 自动从 PDF 提取
    2. 预处理:   剥离 ctexart 版式命令, 图/表/公式编号文本化, \\ref/\\cite 解析
    3. pandoc:   tex -> docx (公式转 Word 原生 OMML, 可编辑)
    4. 后处理:   套模板样式(宋体/黑体/行距/缩进), 公式编号, 三线表, 页脚页码
    5. 验证:     结构统计(段落/表格/图片/公式数)

依赖: pandoc, python-docx, PyMuPDF(fitz)
"""
import sys, os, re, glob, subprocess

HERE = os.path.dirname(os.path.abspath(__file__))

def tex_image_names(tex_path):
    src = open(tex_path, encoding='utf-8').read()
    return re.findall(r'\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}', src)

def extract_images_from_pdf(pdf_path, names, outdir):
    import fitz
    doc = fitz.open(pdf_path)
    imgs = []
    for pno in range(len(doc)):
        for img in doc[pno].get_images(full=True):
            pix = fitz.Pixmap(doc, img[0])
            imgs.append((pno, img[0], pix))
    print(f'[1/4] PDF 内嵌图片 {len(imgs)} 张, 需要 {len(names)} 张')
    if len(imgs) != len(names):
        print('      警告: 图片数量不匹配, 按顺序提取前 %d 张' % len(names))
    for i, (pno, xref, pix) in enumerate(imgs):
        if i >= len(names):
            break
        if pix.n - pix.alpha >= 4:
            pix = fitz.Pixmap(fitz.csRGB, pix)
        out = os.path.join(outdir, names[i])
        pix.save(out)
        print(f'      提取: {names[i]} (p{pno+1})')

def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    tex = os.path.abspath(sys.argv[1])
    workdir = os.path.dirname(tex)
    base = os.path.splitext(os.path.basename(tex))[0]
    ref = os.path.abspath(sys.argv[2]) if len(sys.argv) > 2 else ''
    out_docx = os.path.abspath(sys.argv[3]) if len(sys.argv) > 3 else os.path.join(workdir, base + '.docx')

    os.chdir(workdir)
    clean_tex = os.path.join(workdir, 'paper_clean.tex')
    labels = os.path.join(workdir, 'labels.json')

    if not ref:
        cand = glob.glob(os.path.join(workdir, '*.dotx')) + glob.glob(os.path.join(workdir, '*模板*.docx'))
        if not cand:
            print('错误: 未找到模板 (请传 reference.docx/dotx 参数或放模板在目录)')
            sys.exit(1)
        ref = cand[0]

    # 1. 图片修复
    names = [os.path.basename(n) for n in tex_image_names(tex)]
    missing = [n for n in names if not os.path.exists(os.path.join(workdir, n))]
    if missing:
        pdfs = glob.glob(os.path.join(workdir, '*.pdf'))
        if pdfs:
            print(f'[1/4] 图片缺失 {len(missing)} 张, 从 {os.path.basename(pdfs[0])} 提取')
            extract_images_from_pdf(pdfs[0], names, workdir)
        else:
            print(f'[1/4] 警告: 图片缺失且无 PDF 可提取: {missing}')
    else:
        print('[1/4] 图片齐全')

    # 2. 预处理
    subprocess.run([sys.executable, os.path.join(HERE, 'preprocess.py'), tex, clean_tex, labels], check=True)
    print('[2/4] 预处理完成')

    # 3. pandoc
    subprocess.run(['pandoc', clean_tex, '-f', 'latex', '-t', 'docx',
                    '--reference-doc', ref, '-o', out_docx], check=True)
    print('[3/4] pandoc 转换完成')

    # 4. 后处理
    subprocess.run([sys.executable, os.path.join(HERE, 'postprocess.py'), out_docx], check=True)
    print('[4/4] 后处理完成')

    # 5. 验证
    import docx as _docx
    from docx.oxml.ns import qn
    d = _docx.Document(out_docx)
    body = d.element.body
    print(f'[5/5] 验证: 段落 {len(d.paragraphs)}, 表格 {len(d.tables)}, '
          f'图片 {len(body.findall(".//" + qn("w:drawing")))}, '
          f'独立公式 {len(body.findall(".//" + qn("m:oMathPara")))}')
    print(f'完成: {out_docx}')

if __name__ == '__main__':
    main()
