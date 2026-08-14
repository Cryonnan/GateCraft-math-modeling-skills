# -*- coding: utf-8 -*-
"""后处理：套用数模模板样式、公式编号、三线表、页码、代码/标题/表格修正
用法: python postprocess.py [input.docx]"""
import docx, re, sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = sys.argv[1] if len(sys.argv) > 1 else r'NIPT_paper_v7_final.docx'
d = Document(SRC)
body = d.element.body

SONG = '宋体'
HEI = '黑体'
TNR = 'Times New Roman'

def set_style_font(style, ascii_f, east_f, size=None, bold=None):
    rpr = style.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    rf.set(qn('w:ascii'), ascii_f)
    rf.set(qn('w:hAnsi'), ascii_f)
    rf.set(qn('w:eastAsia'), east_f)
    if size is not None:
        for tag in ('w:sz', 'w:szCs'):
            sz = rpr.find(qn(tag))
            if sz is None:
                sz = OxmlElement(tag); rpr.append(sz)
            sz.set(qn('w:val'), str(int(size * 2)))
    if bold is not None:
        b = rpr.find(qn('w:b'))
        if b is None:
            b = OxmlElement('w:b'); rpr.append(b)
        if not bold:
            b.set(qn('w:val'), '0')

def remove_spacing_override(style):
    """移除样式级段前段后(确保正文段前后0行)"""
    pPr = style.element.find(qn('w:pPr'))
    if pPr is None: return
    for tag in ('w:spacing', 'w:ind'):
        el = pPr.find(qn(tag))
        if el is not None:
            pPr.remove(el)

def add_char_font(run, ascii_f=TNR, east_f=SONG, bold=None, size=None, italic=None):
    r = run._r
    rPr = r.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    rf.set(qn('w:ascii'), ascii_f); rf.set(qn('w:hAnsi'), ascii_f)
    rf.set(qn('w:eastAsia'), east_f)
    if size is not None:
        for tag in ('w:sz', 'w:szCs'):
            sz = rPr.find(qn(tag))
            if sz is None:
                sz = OxmlElement(tag); rPr.append(sz)
            sz.set(qn('w:val'), str(int(size * 2)))
    if bold is not None:
        b = rPr.find(qn('w:b'))
        if b is None:
            b = OxmlElement('w:b'); rPr.append(b)
        if not bold:
            b.set(qn('w:val'), '0')
    if italic is not None:
        it = rPr.find(qn('w:i'))
        if it is None:
            it = OxmlElement('w:i'); rPr.append(it)
        if not italic:
            it.set(qn('w:val'), '0')

# ---------- 1. 样式级：正文/标题/代码 ----------
remove_spacing_patch = [None]
set_style_font(d.styles['Normal'], TNR, SONG, 12, False)
for sname in ('Body Text', 'First Paragraph'):
    try:
        set_style_font(d.styles[sname], TNR, SONG, 12, False)
        remove_spacing_override(d.styles[sname])
    except KeyError:
        pass
# 标题字体：黑体，H1=22pt(二号) H2=16pt(三号) H3=14pt(四号)，模板为准
for sname, sz in (('Heading 1', 22), ('Heading 2', 16), ('Heading 3', 14)):
    try:
        set_style_font(d.styles[sname], TNR, HEI, sz, True)
        pPr = d.styles[sname].element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr'); d.styles[sname].element.insert(0, pPr)
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing'); pPr.append(sp)
        sp.set(qn('w:line'), '240'); sp.set(qn('w:lineRule'), 'auto')
        if sname == 'Heading 1':
            sp.set(qn('w:before'), '240'); sp.set(qn('w:after'), '240')
        elif sname == 'Heading 2':
            sp.set(qn('w:before'), '240'); sp.set(qn('w:after'), '120')
        else:
            sp.set(qn('w:before'), '120'); sp.set(qn('w:after'), '60')
    except KeyError:
        pass
try:
    set_style_font(d.styles['Source Code'], 'Courier New', SONG, 12)
    remove_spacing_override(d.styles['Source Code'])
except KeyError:
    pass
try:
    d.styles['Caption'].font.size = Pt(10.5)
except KeyError:
    pass

# ---------- 2. 段落级处理 ----------
CAP_RE = re.compile(r'^(表|图)\s*\d+(\.\d+)*\s')
cur_sec = 0
appendix = False
eq_cnt = 0
CONTENT_W = 8313

def add_eq_number(p, num):
    pPr = p._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        if len(pPr) and pPr[0].tag == qn('w:pStyle'):
            pPr.insert(1, tabs)
        else:
            pPr.insert(0, tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(CONTENT_W))
    tabs.append(tab)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), TNR); rf.set(qn('w:hAnsi'), TNR); rf.set(qn('w:eastAsia'), SONG)
    rPr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24'); rPr.append(sz)
    r.append(rPr)
    tabr = OxmlElement('w:tab')
    r.append(tabr)
    t = OxmlElement('w:t')
    t.text = num
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p._p.append(r)

def set_para_font(p, east, size_pt, bold=True):
    for run in p.runs:
        add_char_font(run, east_f=east, size=size_pt, bold=bold)

def set_indent_and_spacing(p, indent_chars=200, line=240, before=0, after=0):
    """默认单倍行距(line=240=1.0)，首行缩进2字符，段前段后0"""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(indent_chars))
    ind.set(qn('w:firstLine'), str(int(indent_chars * 60 * 12 / 100)))
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing'); pPr.append(spacing)
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if before:
        spacing.set(qn('w:beforeLines'), str(before))
    if after:
        spacing.set(qn('w:afterLines'), str(after))

def disable_numpr(p):
    """段落级禁用自动编号（numId=0 覆盖样式级 numPr）"""
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn('w:numPr'))
    if old is not None:
        pPr.remove(old)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0'); numPr.append(ilvl)
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '0'); numPr.append(numId)
    pPr.insert(0, numPr)

def add_page_break_before(p):
    pPr = p._p.get_or_add_pPr()
    br = OxmlElement('w:pageBreakBefore')
    pPr.insert(0, br)

first_title_done = False
for p in d.paragraphs:
    st = p.style.name
    txt = p.text.strip()

    if st == 'Heading 1':
        if '参考文献' in txt or '附录' in txt:
            # 参考文献/附录标题：另起一页 + 不编号
            add_page_break_before(p)
            disable_numpr(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            eq_cnt = 0
            if '附录' in txt: appendix = True
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cur_sec += 1
        eq_cnt = 0
        continue

    if 'Appendix' in st or st == 'Heading 3' and appendix:
        continue

    has_display_math = bool(p._element.findall(qn('m:oMathPara')))
    if has_display_math:
        eq_cnt += 1
        num = f'({cur_sec}.{eq_cnt})' if not appendix else f'(A.{eq_cnt})'
        add_eq_number(p, num)
        p.style = d.styles['Normal']
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_indent_and_spacing(p, indent_chars=0)
        continue

    if st in ('Heading 2', 'Heading 3'):
        continue

    if 'Source' in st or st == 'Source Code':
        # 代码段：单倍行距、无首行缩进、等宽字体已由样式保证
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:firstLine'), '0')
        continue

    if CAP_RE.match(txt):
        # 图表标题：居中，黑体五号（caption 样式已居中）
        set_para_font(p, HEI, 10.5, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue

    if not first_title_done and txt.startswith('NIPT'):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_font(p, HEI, 17.28)
        first_title_done = True
        continue

    if txt == '摘 要' or txt.startswith('摘要'):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_font(p, HEI, 14)
        continue

    if txt.startswith('关键词：'):
        set_para_font(p, HEI, 12, True)
        continue

    if st == 'List Paragraph' and txt:
        # 列表项：保持项目符号，单倍行距，不设首行缩进
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind'); pPr.append(ind)
        ind.set(qn('w:firstLineChars'), '0')
        ind.set(qn('w:firstLine'), '0')
        if not ind.get(qn('w:left')):
            ind.set(qn('w:left'), '360')
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing'); pPr.append(spacing)
        spacing.set(qn('w:line'), '240'); spacing.set(qn('w:lineRule'), 'auto')
        spacing.set(qn('w:before'), '0'); spacing.set(qn('w:after'), '0')
        for run in p.runs:
            add_char_font(run, east_f=SONG, size=12, bold=False)
        continue

    if st in ('Body Text', 'First Paragraph', 'Normal') and txt:
        # 正文：单倍行距 + 首行缩进2字符 + 段前段后0 + 样式统一为 Normal(命"正文")
        pPr = p._p.get_or_add_pPr()
        has_num = pPr.find(qn('w:numPr')) is not None
        is_ref = bool(re.match(r'^\[\d+\]', txt))
        set_indent_and_spacing(p, indent_chars=0 if (has_num or is_ref) else 200)
        if has_num:
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind'); pPr.append(ind)
            ind.set(qn('w:left'), '360')
        if is_ref:
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind'); pPr.append(ind)
            ind.set(qn('w:left'), '360')
            ind.set(qn('w:hanging'), '360')
        p.style = d.styles['Normal']

# ---------- 3. 表格 ----------
def tbl_border(el, attrs):
    for k, v in attrs.items():
        el.set(qn('w:' + k), v)

for tbl in d.tables:
    tblPr = tbl._tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for tag, val, sz in (('w:top', 'single', 12), ('w:bottom', 'single', 12),
                         ('w:left', 'nil', 0), ('w:right', 'nil', 0),
                         ('w:insideH', 'nil', 0), ('w:insideV', 'nil', 0)):
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); borders.append(el)
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), '000000')
    # 表头行下细线 + 表头加粗水平居中
    first = tbl.rows[0]
    for cell in first.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tb = tcPr.find(qn('w:tcBorders'))
        if tb is None:
            tb = OxmlElement('w:tcBorders'); tcPr.append(tb)
        bottom = tb.find(qn('w:bottom'))
        if bottom is None:
            bottom = OxmlElement('w:bottom'); tb.append(bottom)
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '000000')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                add_char_font(run, east_f=SONG, size=12, bold=True)
    # 表格内单元格：去除缩进，统一宋体，单倍行距，水平居中
    for row in tbl.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                pPr = para._p.get_or_add_pPr()
                ind = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = OxmlElement('w:ind'); pPr.append(ind)
                ind.set(qn('w:firstLineChars'), '0')
                ind.set(qn('w:firstLine'), '0')
                ind.set(qn('w:left'), '0')
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing'); pPr.append(spacing)
                spacing.set(qn('w:line'), '240')
                spacing.set(qn('w:lineRule'), 'auto')
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
                for run in para.runs:
                    add_char_font(run, east_f=SONG, size=12, bold=False)

# ---------- 4. 页脚页码 ----------
sec = d.sections[0]
footer = sec.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
rPr = run._r.get_or_add_rPr()
rf = OxmlElement('w:rFonts')
rf.set(qn('w:ascii'), TNR); rf.set(qn('w:eastAsia'), SONG)
rPr.insert(0, rf)
sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18'); rPr.append(sz)
fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

d.save(SRC)
print('postprocess done')